#!/usr/bin/env python3
"""
Apply justified text alignment to all paragraphs in a DOCX file.
Preserves images, formatting, and structure.
"""

import sys
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def justify_docx(input_path, output_path):
    """Load DOCX, apply justified alignment to all paragraphs, save to output."""
    doc = Document(input_path)

    # Apply justified alignment to all paragraphs
    for paragraph in doc.paragraphs:
        # Don't justify headings (keep them left-aligned)
        if paragraph.style.name.startswith('Heading'):
            continue

        # Apply justified alignment to body text
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Save modified document
    doc.save(output_path)
    print(f"✓ Justified text applied: {output_path}")
    print(f"  Total paragraphs processed: {len(doc.paragraphs)}")

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python justify-docx.py <input.docx> <output.docx>")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]

    justify_docx(input_file, output_file)
