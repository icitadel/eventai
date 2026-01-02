"""Core functionality for DOCX document analysis and formatting."""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import re
import subprocess
from pathlib import Path
from typing import Dict, Tuple, Optional


def word_count(filepath: str) -> int:
    """
    Count words in a DOCX file, excluding headers and footers.
    
    Args:
        filepath: Path to the DOCX file
        
    Returns:
        Total word count
    """
    doc = Document(filepath)
    
    total_words = 0
    for paragraph in doc.paragraphs:
        # Count words in paragraph text
        text = paragraph.text.strip()
        if text:
            words = len(text.split())
            total_words += words
    
    return total_words


def heading_count(filepath: str) -> Dict[str, int]:
    """
    Count headings by level in a DOCX file.
    
    Args:
        filepath: Path to the DOCX file
        
    Returns:
        Dictionary mapping heading levels to counts (e.g., {'H1': 5, 'H2': 12})
    """
    doc = Document(filepath)
    
    heading_counts = {}
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name
        if style_name.startswith('Heading'):
            # Extract heading level (e.g., 'Heading 1' -> 'H1')
            match = re.search(r'Heading (\d)', style_name)
            if match:
                level = f"H{match.group(1)}"
                heading_counts[level] = heading_counts.get(level, 0) + 1
    
    return heading_counts


def density_metrics(filepath: str) -> Dict[str, float]:
    """
    Calculate density metrics for a DOCX file.
    
    Metrics calculated:
    - words_per_sentence: Average words per sentence
    - sentences_per_paragraph: Average sentences per paragraph
    - paragraphs_per_section: Average paragraphs per section (between headings)
    
    Args:
        filepath: Path to the DOCX file
        
    Returns:
        Dictionary of density metrics
    """
    doc = Document(filepath)
    
    total_words = 0
    total_sentences = 0
    total_paragraphs = 0
    section_count = 0
    paragraphs_in_current_section = 0
    total_paragraphs_across_sections = 0
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        
        # Check if this is a heading (new section)
        if paragraph.style.name.startswith('Heading'):
            if paragraphs_in_current_section > 0:
                total_paragraphs_across_sections += paragraphs_in_current_section
                section_count += 1
            paragraphs_in_current_section = 0
            continue
        
        if not text:
            continue
        
        # Count paragraphs (non-empty, non-heading)
        total_paragraphs += 1
        paragraphs_in_current_section += 1
        
        # Count words
        words = text.split()
        total_words += len(words)
        
        # Count sentences (simple heuristic: split on . ! ?)
        sentences = re.split(r'[.!?]+', text)
        sentences = [s for s in sentences if s.strip()]
        total_sentences += len(sentences)
    
    # Handle last section
    if paragraphs_in_current_section > 0:
        total_paragraphs_across_sections += paragraphs_in_current_section
        section_count += 1
    
    # Calculate metrics
    metrics = {}
    metrics['words_per_sentence'] = total_words / total_sentences if total_sentences > 0 else 0
    metrics['sentences_per_paragraph'] = total_sentences / total_paragraphs if total_paragraphs > 0 else 0
    metrics['paragraphs_per_section'] = total_paragraphs_across_sections / section_count if section_count > 0 else 0
    
    # Add raw counts for context
    metrics['total_words'] = total_words
    metrics['total_sentences'] = total_sentences
    metrics['total_paragraphs'] = total_paragraphs
    metrics['total_sections'] = section_count
    
    return metrics


def resize_images(filepath: str, width_inches: float = 6.5, wrap_style: str = "top-bottom") -> int:
    """
    Resize all images in a DOCX file to specified width, maintaining aspect ratio.
    Also sets text wrapping style and centers images.
    
    Args:
        filepath: Path to the DOCX file
        width_inches: Target width in inches (default: 6.5)
        wrap_style: Text wrapping style - "top-bottom" (Above and Below) or "inline" (default: "top-bottom")
        
    Returns:
        Number of images processed
    """
    doc = Document(filepath)
    
    # Resize all inline shapes
    for shape in doc.inline_shapes:
        # Calculate new height to maintain aspect ratio
        aspect_ratio = shape.height / shape.width
        new_width = Inches(width_inches)
        new_height = int(new_width * aspect_ratio)
        
        shape.width = new_width
        shape.height = new_height
    
    if wrap_style == "top-bottom":
        # Set text wrapping to "Top and Bottom" (Above and Below) for all images
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                # Find drawing elements (images)
                for drawing in run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                    # Check for existing anchor elements
                    anchor = drawing.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor')
                    if anchor is not None:
                        # Remove any existing wrap elements
                        for wrap_elem in ['wrapNone', 'wrapSquare', 'wrapTight', 'wrapThrough']:
                            existing = anchor.find(f'.//{{{nsdecls("wp")}}}{wrap_elem}')
                            if existing is not None:
                                anchor.remove(existing)
                        
                        # Add wrapTopAndBottom element if not present
                        wrap_top_bottom = anchor.find(f'.//{{{nsdecls("wp")}}}wrapTopAndBottom')
                        if wrap_top_bottom is None:
                            # Insert wrapTopAndBottom after effectExtent
                            effect_extent = anchor.find(f'.//{{{nsdecls("wp")}}}effectExtent')
                            if effect_extent is not None:
                                idx = list(anchor).index(effect_extent) + 1
                                wrap_elem = parse_xml(f'<wp:wrapTopAndBottom xmlns:wp="{nsdecls("wp")}"/>')
                                anchor.insert(idx, wrap_elem)
            
            # Center align paragraphs containing images
            has_image = any(run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing') 
                           for run in paragraph.runs)
            if has_image:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.save(filepath)
    return len(doc.inline_shapes)


def justify_text(filepath: str) -> int:
    """
    Apply justified alignment to body paragraphs in a DOCX file.
    Headings and special styles remain left-aligned.

    Args:
        filepath: Path to the DOCX file

    Returns:
        Number of paragraphs justified
    """
    doc = Document(filepath)

    modified_count = 0
    for paragraph in doc.paragraphs:
        # Skip headings (keep left-aligned)
        if paragraph.style.name.startswith('Heading'):
            continue

        # Skip blockquotes and other special styles
        if paragraph.style.name in ['Quote', 'Block Text', 'Caption']:
            continue

        # Apply justified alignment to body text
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        modified_count += 1

    doc.save(filepath)
    return modified_count


def convert_to_docx(input_filepath: str, output_filepath: Optional[str] = None,
                    resource_path: Optional[str] = None,
                    image_width: Optional[float] = None,
                    justify: bool = False) -> str:
    """
    Convert a file to DOCX format using Pandoc, with optional post-processing.

    Supported input formats: Markdown (.md), HTML (.html), reStructuredText (.rst),
    LaTeX (.tex), and any other format Pandoc supports.

    Args:
        input_filepath: Path to the input file
        output_filepath: Path for output DOCX file (default: replace input extension with .docx)
        resource_path: Path for resolving relative paths (e.g., images)
        image_width: If provided, resize all images to this width in inches (default: None)
        justify: If True, apply justified alignment to body paragraphs (default: False)

    Returns:
        Path to the generated DOCX file
    """
    input_path = Path(input_filepath)

    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_filepath}")

    # Determine output filepath
    if output_filepath is None:
        output_filepath = str(input_path.with_suffix('.docx'))

    # Build pandoc command
    cmd = ['pandoc', str(input_filepath), '-o', output_filepath]

    # Add resource path if provided
    if resource_path:
        cmd.extend(['--resource-path', resource_path])

    # Add memory limits for large files
    cmd.extend(['-M2GB', '+RTS', '-M2G', '-RTS'])

    try:
        result = subprocess.run(cmd, check=True, capture_output=True, text=True)
    except subprocess.CalledProcessError as e:
        raise RuntimeError(f"Pandoc conversion failed: {e.stderr}") from e
    except FileNotFoundError:
        raise RuntimeError("Pandoc not found. Please install pandoc: brew install pandoc")

    # Post-processing: resize images if requested
    if image_width is not None:
        resize_images(output_filepath, width_inches=image_width, wrap_style='top-bottom')

    # Post-processing: justify text if requested
    if justify:
        justify_text(output_filepath)

    return output_filepath
